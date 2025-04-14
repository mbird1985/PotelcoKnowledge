import sqlite3
import os
import PyPDF2
from config import SCHEDULE_DB, DOCUMENT_FOLDER, ES_HOST, ES_USER, ES_PASS
from werkzeug.utils import secure_filename
from pdfminer.high_level import extract_text as pdf_extract_text
from elasticsearch import Elasticsearch
from docx import Document
from services.logging_service import log_audit

ALLOWED_EXTENSIONS = {"pdf", "docx", "txt"}

es = Elasticsearch([ES_HOST], http_auth=(ES_USER, ES_PASS))

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def save_uploaded_document(file, user_id):
    filename = secure_filename(file.filename)
    filepath = os.path.join(DOCUMENT_FOLDER, filename)
    file.save(filepath)
    conn = sqlite3.connect(SCHEDULE_DB)
    c = conn.cursor()
    c.execute("INSERT INTO documents (filename, uploader_id) VALUES (?, ?)", (filename, user_id))
    conn.commit()
    conn.close()
    return filename

def get_all_documents():
    conn = sqlite3.connect(SCHEDULE_DB)
    c = conn.cursor()
    c.execute("SELECT * FROM documents ORDER BY id DESC")
    docs = c.fetchall()
    conn.close()
    return docs

def get_document_by_id(doc_id):
    conn = sqlite3.connect(SCHEDULE_DB)
    c = conn.cursor()
    c.execute("SELECT * FROM documents WHERE id = ?", (doc_id,))
    doc = c.fetchone()
    conn.close()
    return doc

def update_document(doc_id, new_title, user_id):
    conn = sqlite3.connect(SCHEDULE_DB)
    c = conn.cursor()
    c.execute("UPDATE documents SET filename = ?, uploader_id = ? WHERE id = ?", (new_title, user_id, doc_id))
    conn.commit()
    conn.close()

def delete_document(doc_id):
    conn = sqlite3.connect(SCHEDULE_DB)
    c = conn.cursor()
    c.execute("SELECT filename FROM documents WHERE id = ?", (doc_id,))
    result = c.fetchone()
    if result:
        file_path = os.path.join(DOCUMENT_FOLDER, result[0])
        if os.path.exists(file_path):
            os.remove(file_path)
    c.execute("DELETE FROM documents WHERE id = ?", (doc_id,))
    conn.commit()
    conn.close()

def get_document_versions(title):
    conn = sqlite3.connect(SCHEDULE_DB)
    c = conn.cursor()
    c.execute("SELECT * FROM documents WHERE filename LIKE ?", (f"{title}%",))
    versions = c.fetchall()
    conn.close()
    return versions

def extract_text(filepath):
    if filepath.endswith(".pdf"):
        return pdf_extract_text(filepath)
    elif filepath.endswith(".docx"):
        doc = docx.Document(filepath)
        return "\n".join([p.text for p in doc.paragraphs])
    else:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()
        
def index_document(file_path, filename, user_id):
    try:
        if filename.endswith('.pdf'):
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = ''.join([page.extract_text() for page in reader.pages])
        elif filename.endswith('.docx'):
            doc = Document(file_path)
            text = '\n'.join([para.text for para in doc.paragraphs])
        else:
            raise ValueError("Unsupported file type")

        chunks = [text[i:i+1000] for i in range(0, len(text), 1000)]
        for i, chunk in enumerate(chunks):
            es.index(index='knowledge_base', id=f"{filename}_{i}", body={'text': chunk, 'filename': filename})
        log_audit(user_id, "index_document", {"filename": filename})
    except Exception as e:
        log_audit(user_id, "index_document_error", {"filename": filename, "error": str(e)})
        raise  # Re-raise to inform the caller (e.g., the route)
